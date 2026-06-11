"""Export Service - SAE report export to Word/PDF/JSON"""
import json
import os
from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── PDF 渲染（fpdf2）──
_FONT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "assets", "fonts")
_FONT_CUSTOM_PATH = os.path.join(_FONT_DIR, "NotoSansSC-Regular.ttf")


def _find_cjk_font():
    """查找可用的中文字体，返回 (font_path, font_name) 或 (None, None)"""
    # 优先级 1: 项目内嵌字体
    if os.path.exists(_FONT_CUSTOM_PATH):
        return _FONT_CUSTOM_PATH, "NotoSansSC"
    # 优先级 2: Windows 系统字体
    win_fonts = [
        (r"C:\Windows\Fonts\msyh.ttc", "msyh"),
        (r"C:\Windows\Fonts\simsun.ttc", "simsun"),
        (r"C:\Windows\Fonts\simhei.ttf", "simhei"),
    ]
    for path, name in win_fonts:
        if os.path.exists(path):
            return path, name
    # 优先级 3: Linux / macOS
    nix_fonts = [
        ("/usr/share/fonts/truetype/noto/NotoSansSC-Regular.ttf", "NotoSansSC"),
        ("/System/Library/Fonts/PingFang.ttc", "PingFang"),
    ]
    for path, name in nix_fonts:
        if os.path.exists(path):
            return path, name
    return None, None


def _build_pdf(report: dict) -> BytesIO:
    """使用 fpdf2 构建 CIOMS-I PDF"""
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=20)

    # 加载中文字体（如果可用）
    font_path, font_name = _find_cjk_font()
    if font_path:
        try:
            font_style = "B" if "Bold" in font_path or "bd" in font_path.lower() else ""
            pdf.add_font(font_name, font_style, font_path)
            pdf.add_font(font_name, "B", font_path)  # 粗体回退
        except Exception:
            font_name = "Helvetica"
    else:
        font_name = "Helvetica"
    font_name = "Helvetica"
    if font_ok:
        try:
            pdf.add_font("NotoSansSC", "", _FONT_PATH, uni=True)
            pdf.add_font("NotoSansSC", "B", _FONT_BOLD_PATH, uni=True)
            font_name = "NotoSansSC"
        except Exception:
            font_name = "Helvetica"

    # Helper functions
    def set_font(style="", size=10):
        pdf.set_font(font_name, style, size)

    def section(title, fields):
        set_font("B", 13)
        pdf.cell(0, 10, title, ln=True)
        for label, value in fields:
            set_font("", 10)
            val = str(value or "")
            text = f"{label}: {val}"
            # multi_cell handles long text
            pdf.multi_cell(0, 6, text)
        pdf.ln(3)

    def section_text(title, body):
        set_font("B", 13)
        pdf.cell(0, 10, title, ln=True)
        set_font("", 10)
        pdf.multi_cell(0, 6, str(body or ""))
        pdf.ln(3)

    # ═════ Build PDF ═════
    pdf.add_page()

    # 标题
    set_font("B", 18)
    pdf.cell(0, 15, "CIOMS-I", ln=True, align="C")
    set_font("B", 14)
    pdf.cell(0, 10, "严重不良事件报告", ln=True, align="C")
    pdf.ln(5)

    cioms = report.get("cioms_fields", {})

    # 报告元信息
    set_font("", 9)
    pdf.cell(0, 5, f"报告编号: {report.get('report_id', '')}", ln=True)
    pdf.cell(0, 5, f"报告状态: {report.get('report_status', 'draft')}", ln=True)
    pdf.cell(0, 5, f"生成日期: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ln=True)
    pdf.ln(5)

    # 各章节
    section("1. 患者信息", [
        ("患者编号", cioms.get("patient_initials")),
        ("性别", cioms.get("patient_gender")),
        ("出生日期", cioms.get("patient_dob")),
    ])

    section("2. 不良事件", [
        ("AE 描述", cioms.get("ae_description")),
        ("开始日期", cioms.get("ae_start_date")),
        ("结束日期", cioms.get("ae_end_date")),
        ("严重性", "是" if cioms.get("ae_serious") else "否"),
        ("转归", cioms.get("ae_outcome")),
    ])

    section("3. 可疑药物", [
        ("药物名称", cioms.get("suspect_drug_name")),
        ("剂量", cioms.get("suspect_drug_dose")),
        ("用药日期", cioms.get("suspect_drug_dates")),
        ("合并用药", cioms.get("concomitant_drugs")),
    ])

    section("4. 因果关系评估", [
        ("评估方法", cioms.get("causality_method")),
        ("评分", cioms.get("causality_score") or report.get("causality_assessment", "")),
        ("去激发", cioms.get("dechallenge")),
        ("再激发", cioms.get("rechallenge")),
    ])

    section("5. 报告人", [
        ("资质", cioms.get("reporter_qualification")),
        ("研究编号", cioms.get("study_number")),
        ("申办方", cioms.get("sponsor")),
    ])

    section_text("6. 事件叙述", cioms.get("narrative"))
    section_text("7. 同类药物安全性信息", report.get("similar_drug_safety"))

    buffer = BytesIO()
    pdf.output(buffer)
    buffer.seek(0)
    return buffer


# ── 导出函数 ──

def export_to_docx(report: dict) -> BytesIO:
    doc = Document()
    title = doc.add_heading('CIOMS-I 严重不良事件报告', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"报告编号: {report.get('report_id', '')}")
    doc.add_paragraph(f"研究编号: {report.get('cioms_fields', {}).get('study_number', '')}")
    doc.add_paragraph(f"生成日期: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph(f"报告状态: {report.get('report_status', 'draft')}")
    doc.add_paragraph("")
    cioms = report.get('cioms_fields', {})
    sections = [
        ("患者信息", [("患者编号", cioms.get("patient_initials", "")), ("性别", cioms.get("patient_gender", "")), ("出生日期", cioms.get("patient_dob", ""))]),
        ("不良事件信息", [("AE描述", cioms.get("ae_description", "")), ("开始日期", cioms.get("ae_start_date", "")), ("结束日期", cioms.get("ae_end_date", "")), ("是否严重", "是" if cioms.get("ae_serious") else "否"), ("转归", cioms.get("ae_outcome", ""))]),
        ("可疑药物信息", [("药物名称", cioms.get("suspect_drug_name", "")), ("剂量", cioms.get("suspect_drug_dose", "")), ("用药日期", cioms.get("suspect_drug_dates", ""))]),
        ("因果关系评估", [("评估方法", cioms.get("causality_method", "")), ("评估结果", report.get('causality_assessment', '')), ("去激发", cioms.get("dechallenge", "")), ("再激发", cioms.get("rechallenge", ""))]),
        ("报告人信息", [("资质", cioms.get("reporter_qualification", ""))]),
    ]
    for section_title, fields in sections:
        doc.add_heading(section_title, level=2)
        for label, value in fields:
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            run_label.font.size = Pt(11)
            run_value = p.add_run(str(value))
            run_value.font.size = Pt(11)
    doc.add_heading("事件叙述", level=2)
    doc.add_paragraph(cioms.get("narrative", ""))
    doc.add_heading("同类药物安全性信息", level=2)
    doc.add_paragraph(report.get('similar_drug_safety', ''))
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def export_to_json(report: dict) -> BytesIO:
    buffer = BytesIO()
    json_bytes = json.dumps(report, ensure_ascii=False, indent=2).encode('utf-8')
    buffer.write(json_bytes)
    buffer.seek(0)
    return buffer


def export_to_pdf(report: dict) -> BytesIO:
    """生成 CIOMS-I PDF 报告，真正 PDF 格式（fpdf2）"""
    return _build_pdf(report)
